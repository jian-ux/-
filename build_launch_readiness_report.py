from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\workspace\feisheng-bot\docs\Feisheng-Bot上线准备度评估报告.docx")

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "20262E"
WHITE = "FFFFFF"
RED = "9B1C1C"
RED_FILL = "FDECEC"
AMBER = "7A5A00"
AMBER_FILL = "FFF7D6"
GREEN = "2F6B4F"
GREEN_FILL = "EAF5EF"
BORDER = "D0D5DD"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=10.5, bold=False, color=DARK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, size=10.5, bold=False, color=DARK, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return run


def style_paragraph(paragraph, before=0, after=6, line=1.10, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_together = keep
    return paragraph


def add_body(doc, text, bold_lead=None, after=6, keep=False):
    p = doc.add_paragraph()
    style_paragraph(p, after=after, line=1.10, keep=keep)
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, bold=True)
        add_text(p, text[len(bold_lead):])
    else:
        add_text(p, text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
    else:
        set_run_font(run, size=11.5, bold=True, color=DARK_BLUE)
    return p


def add_bullet(doc, text, level=0, color=DARK, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    style_paragraph(p, after=4, line=1.10)
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, bold=True, color=color)
        add_text(p, text[len(bold_lead):], color=color)
    else:
        add_text(p, text, color=color)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, after=5, line=1.10)
    add_text(p, text)
    return p


def add_callout(doc, label, text, severity="info"):
    colors = {
        "risk": (RED_FILL, RED),
        "warn": (AMBER_FILL, AMBER),
        "good": (GREEN_FILL, GREEN),
        "info": (LIGHT_BLUE, DARK_BLUE),
    }
    fill, accent = colors[severity]
    p = doc.add_paragraph()
    style_paragraph(p, before=4, after=10, line=1.10, keep=True)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.12)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "5")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    add_text(p, f"{label}  ", size=10.5, bold=True, color=accent)
    add_text(p, text, size=10.5, color=DARK)
    return p


def add_status_table(doc, headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 or idx == status_col else WD_ALIGN_PARAGRAPH.LEFT
        style_paragraph(p, after=0, line=1.0)
        add_text(p, header, size=9.5, bold=True, color=NAVY)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            if status_col is not None and idx == status_col:
                status = str(value)
                if "阻断" in status or "P0" in status or "高" in status:
                    set_cell_shading(cell, RED_FILL)
                    color = RED
                elif "重要" in status or "P1" in status or "中" in status:
                    set_cell_shading(cell, AMBER_FILL)
                    color = AMBER
                elif "通过" in status or "良好" in status or "低" in status:
                    set_cell_shading(cell, GREEN_FILL)
                    color = GREEN
                else:
                    color = DARK
            else:
                color = DARK
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 or idx == status_col else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.05)
            add_text(p, str(value), size=9.2, bold=(idx == 0 or idx == status_col), color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_issue(doc, number, title, impact, evidence, remediation, acceptance):
    add_heading(doc, f"{number}. {title}", 2)
    p = doc.add_paragraph()
    style_paragraph(p, after=4, line=1.10)
    add_text(p, "上线影响：", bold=True, color=RED)
    add_text(p, impact)
    p = doc.add_paragraph()
    style_paragraph(p, after=4, line=1.10)
    add_text(p, "审计证据：", bold=True, color=DARK_BLUE)
    add_text(p, evidence)
    p = doc.add_paragraph()
    style_paragraph(p, after=4, line=1.10)
    add_text(p, "整改要求：", bold=True, color=AMBER)
    add_text(p, remediation)
    p = doc.add_paragraph()
    style_paragraph(p, after=9, line=1.10)
    add_text(p, "验收标准：", bold=True, color=GREEN)
    add_text(p, acceptance)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.50)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)
    styles["List Bullet 2"].paragraph_format.left_indent = Inches(0.75)
    styles["List Bullet 2"].paragraph_format.first_line_indent = Inches(-0.25)
    styles["List Number"].paragraph_format.left_indent = Inches(0.50)
    styles["List Number"].paragraph_format.first_line_indent = Inches(-0.25)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MID_GRAY)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(hp, after=0, line=1.0)
    add_text(hp, "FEISHENG BOT  |  上线准备度评估", size=8.5, bold=True, color=MID_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(fp, after=0, line=1.0)
    add_text(fp, "内部评审  |  ", size=8.5, color=MID_GRAY)
    add_field(fp, "PAGE")
    add_text(fp, " / ", size=8.5, color=MID_GRAY)
    add_field(fp, "NUMPAGES")


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "Feisheng Bot 上线准备度评估报告"
    doc.core_properties.subject = "生产上线差距、风险及整改路线"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "Feisheng Bot, 上线评估, 风险审计, 整改计划"

    # Opening block: memo_masthead pattern without decorative border.
    p = doc.add_paragraph()
    style_paragraph(p, before=12, after=4, line=1.0)
    add_text(p, "上线准备度评估报告", size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    style_paragraph(p, after=5, line=1.0, keep=True)
    add_text(p, "Feisheng Bot", size=27, bold=True, color=NAVY)

    p = doc.add_paragraph()
    style_paragraph(p, after=16, line=1.05)
    add_text(p, "生产上线差距、风险分级与整改路线", size=14, color=MID_GRAY)

    metadata = [
        ("评估日期", "2026-07-13"),
        ("评估对象", "本地 Docker Compose 部署与当前工作区代码"),
        ("评估口径", "对外生产上线标准，不以演示可用为通过条件"),
        ("文档状态", "内部评审稿"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        style_paragraph(p, after=2, line=1.0)
        add_text(p, f"{label}：", size=9.5, bold=True, color=NAVY)
        add_text(p, value, size=9.5, color=DARK)

    add_callout(
        doc,
        "总体结论",
        "当前版本适合内部演示和小范围联调，但尚未达到对外生产上线标准。主要阻断项集中在权限控制、渠道入口、业务闭环、知识库向量化、数据库迁移和密钥安全。",
        "risk",
    )

    add_heading(doc, "管理层摘要", 1)
    add_body(doc, "基础设施与核心技术链路已经具备可运行基础：MySQL、Redis、MinIO、前后端容器均处于健康状态，管理员登录、真实 LLM 调用、文档存储和基础 FAQ 流程可以工作。")
    add_body(doc, "但“服务可启动”不等于“功能可上线”。当前至少有 7 项 P0 级阻断问题，任何一项未关闭都可能导致越权访问、渠道不可用、配置不生效、知识库检索失真或发布后数据库故障。")

    add_status_table(
        doc,
        ["维度", "当前判断", "核心依据", "状态"],
        [
            ["基础设施", "可运行", "5 个容器健康，MySQL/Redis/MinIO 可用", "通过"],
            ["身份认证", "部分可用", "JWT 登录正常，但权限边界不完整", "P0 阻断"],
            ["渠道接入", "部分可用", "钉钉/企微密钥已注入，网页渠道入口不可用", "P0 阻断"],
            ["知识库", "部分可用", "上传修复完成，但无 Embedding 模型和向量数据", "P0 阻断"],
            ["业务闭环", "不足", "多项页面仅 CRUD，未进入消息处理链路", "P0 阻断"],
            ["发布治理", "不足", "无正式迁移机制，镜像构建跳过测试", "P0 阻断"],
            ["依赖安全", "当前良好", "npm 官方审计未发现已知漏洞", "通过"],
        ],
        [1120, 1650, 4310, 1800],
        status_col=3,
    )

    add_heading(doc, "已确认的积极基础", 1)
    add_bullet(doc, "Docker 运行态正常，前端、后端、MySQL、Redis、MinIO 均为 healthy。")
    add_bullet(doc, "管理员登录和受保护接口可正常工作，真实 LLM 试聊返回 source=ai，抽测延迟约 3.2 秒。")
    add_bullet(doc, "钉钉与企微所需密钥均已注入运行容器；钉钉签名与加解密已有测试覆盖。")
    add_bullet(doc, "后端现有 13 个测试全部通过；前端生产依赖经 npm 官方源审计未发现已知漏洞。")
    add_bullet(doc, "知识库文件可写入 MinIO，并能完成解析、切片及删除回滚。")

    doc.add_page_break()
    add_heading(doc, "P0 上线阻断问题", 1)
    add_callout(doc, "发布门禁", "以下问题必须在生产发布前全部关闭，不建议以“上线后再修”的方式接受风险。", "risk")

    add_issue(
        doc, "P0-1", "RBAC 权限没有完整生效",
        "普通登录用户可以访问知识库、客户、日志、文档和工单等敏感后台资源，存在数据越权与误操作风险。",
        "运行态使用无角色临时账号抽测：用户和角色接口返回 401，但知识库、客户、日志、文档、工单接口均返回 200。SecurityConfig 对大量接口仅要求 authenticated。",
        "建立权限标识到 API 的明确映射，启用方法级鉴权或完整请求匹配；前端菜单和路由同步按权限过滤；增加角色矩阵集成测试。",
        "无角色账号只能访问本人信息；客服、运营和管理员的每项读写权限均有正向与反向自动化测试，越权统一返回 403。",
    )
    add_issue(
        doc, "P0-2", "超级管理员账号保护不完整",
        "最后一个超级管理员仍可能被删除、禁用或解除角色，导致整个后台失去管理入口。",
        "当前仅保护 admin 角色不可删除；UserController 仍允许删除任意用户，并可先删除全部用户角色再重新写入。",
        "禁止删除、禁用或解除最后一个超级管理员；用户删除必须同步清理关联表；角色分配操作使用事务并校验目标角色。",
        "数据库中始终至少存在一个启用的超级管理员；所有破坏最后管理员的请求均被拒绝并写入审计日志。",
    )
    add_issue(
        doc, "P0-3", "网页客服渠道当前不可用",
        "官网聊天组件无法通过当前生产入口发送消息，直接影响核心客服场景。",
        "实测 http://localhost/gateway/channel/web/message 返回 405，直连后端 8082 返回 401；Nginx 未代理 /gateway，Spring Security 也未对网页渠道设计合法放行方案。",
        "增加 /gateway 代理；为网页渠道设计站点密钥、签名、来源白名单、限流和防重放机制；后端端口不直接暴露公网。",
        "通过 HTTPS 域名从真实网页组件完成消息发送、回复、重复消息去重和限流测试，未授权来源被拒绝。",
    )
    add_issue(
        doc, "P0-4", "多项后台功能只有 CRUD，没有业务闭环",
        "管理人员在页面中保存配置后，机器人行为不发生变化，容易造成“配置已生效”的错误认知。",
        "渠道配置、意图管理、回复策略仅由 Controller/Mapper 读写；消息处理主链路没有引用这些实体。客户、工单和操作日志也缺少自动建档、建单和审计写入。",
        "逐项决定“接入主链路”或“暂时下线入口”。已保留功能必须定义触发条件、优先级、异常策略和可验证结果。",
        "每个后台配置都有端到端测试证明保存后会改变机器人行为；未完成的菜单不出现在生产版本。",
    )
    add_issue(
        doc, "P0-5", "文档知识库尚未形成完整 RAG 能力",
        "上传页面可能显示处理完成，但文档没有向量，语义检索无法命中，真实回答质量与后台状态不一致。",
        "当前启用的 Embedding 模型为 0，数据库中带向量的切片为 0；文档即使 embeddedCount=0 仍被标记为 STATUS_COMPLETED。",
        "配置并验证 Embedding 模型；拆分解析、切片、向量化和审核状态；零向量视为失败或降级状态，并允许重试。",
        "至少使用一组业务文档完成上传、向量化、审核、检索、引用和删除回归；后台可区分各处理阶段。",
    )
    add_issue(
        doc, "P0-6", "数据库升级没有正式迁移机制",
        "代码升级后可能因字段缺失直接出现 500，发布不可预测且难以回滚。",
        "本次文档上传故障即由旧 MySQL 数据卷未执行 05_add_minio_storage.sql 引起。docker-entrypoint-initdb.d 仅在首次初始化时运行。",
        "接入 Flyway 或 Liquibase；迁移脚本必须有版本、校验和、前向兼容与失败回滚策略；启动前执行 schema 校验。",
        "从任意受支持旧版本升级到新版本可自动完成迁移；重复启动不会重复执行；迁移失败时应用不得接收流量。",
    )
    add_issue(
        doc, "P0-7", "AI 密钥明文存储且存在日志泄露风险",
        "数据库或日志泄露会直接暴露第三方模型凭据，可能造成费用损失和数据泄露。",
        "BotAiModelConfig 直接保存 apiKey；EncryptionUtil 未被业务调用，且使用固定默认密钥和 AES-ECB；生产配置启用了 MyBatis StdOut SQL 参数日志。",
        "使用 KMS/Secrets Manager 或 AES-GCM 信封加密；移除默认密钥；接口永不回传密钥；生产关闭 SQL 参数日志并轮换现有密钥。",
        "数据库中不可直接识别原始密钥；日志扫描无密钥特征；密钥轮换、解密失败和权限访问均有测试与告警。",
    )

    doc.add_page_break()
    add_heading(doc, "P1 重要问题", 1)
    add_status_table(
        doc,
        ["编号", "问题", "风险", "建议验收结果"],
        [
            ["P1-1", "上传容量配置不一致", "Nginx 默认约 1 MB，实测 2 MB 返回 413；后端声明 50 MB。", "代理层与应用层限制一致，并覆盖 0 B、上限内和超限文件。"],
            ["P1-2", "异步文档任务不可恢复", "任务仅存在 JVM 内存线程池，重启后 processing 记录会永久悬挂。", "使用持久队列/任务表，支持幂等、重试、超时和启动恢复。"],
            ["P1-3", "转人工仅修改状态", "没有客服分配、通知、接单、排队和超时升级。", "完成转接全链路及 SLA 告警，能够追踪处理人和结果。"],
            ["P1-4", "缺少统一限流", "登录、AI、上传和公开回调可能被暴力尝试或刷量。", "按 IP、账号、渠道和租户设置限流，触发后可观测。"],
            ["P1-5", "LLM/Embedding 超时未真正应用", "配置字段存在，但实际使用默认 RestTemplate，外部服务卡住会占满线程。", "连接、读取、总时限及重试策略有自动化故障测试。"],
            ["P1-6", "健康检查过于乐观", "接口固定返回 UP，无法发现 MySQL、Redis、MinIO 异常。", "区分 liveness/readiness，并验证关键依赖和迁移状态。"],
            ["P1-7", "网络边界不适合公网", "仅 HTTP，8082 绑定所有网卡，缺少安全响应头和正式域名配置。", "HTTPS、HSTS、CSP、安全头及仅内网后端端口。"],
            ["P1-8", "CORS 过宽", "允许任意 Origin 且允许凭据，扩大浏览器侧攻击面。", "生产仅允许明确域名，预检和凭据策略经过验证。"],
        ],
        [850, 1900, 3320, 2810],
        status_col=None,
    )

    add_heading(doc, "质量与发布治理差距", 1)
    add_bullet(doc, "Admin 模块没有自动化测试，鉴权、用户、角色、上传、模型配置和迁移缺少覆盖。")
    add_bullet(doc, "前端 package.json 只有 dev/build/preview，没有 lint、unit test 或 E2E 脚本。")
    add_bullet(doc, "生产 Dockerfile 使用 -DskipTests，镜像构建不会因测试失败而停止。")
    add_bullet(doc, "当前工作区存在约 47 项未提交或未跟踪变更，不具备明确版本号、变更集和一键回滚基线。")
    add_bullet(doc, "MyBatis 在各模块启用 StdOutImpl，生产日志噪声大，并可能输出敏感参数。")

    add_heading(doc, "验证结果", 1)
    add_status_table(
        doc,
        ["验证项", "结果", "说明"],
        [
            ["容器状态", "通过", "frontend、app、mysql、redis、minio 均为 healthy"],
            ["管理员登录", "通过", "登录和主要管理员接口返回 200"],
            ["真实 LLM 调用", "通过", "试聊 source=ai，抽测约 3.2 秒"],
            ["后端测试", "通过", "4 个测试类，13 个测试全部通过"],
            ["Admin 自动化测试", "缺失", "Admin 模块 No tests to run"],
            ["前端依赖审计", "通过", "npm 官方源返回 0 vulnerabilities"],
            ["2 MB 文档上传", "失败", "Nginx 返回 413"],
            ["普通账号越权抽测", "失败", "多个敏感后台接口返回 200"],
            ["网页渠道入口", "失败", "代理入口 405，后端直连 401"],
        ],
        [2400, 1500, 4980],
        status_col=1,
    )

    doc.add_page_break()
    add_heading(doc, "整改路线图", 1)
    add_callout(doc, "建议策略", "先收缩生产范围，再补核心闭环。无法在计划周期内完成的功能应隐藏菜单并明确标记为未发布，而不是带着半成品入口上线。", "info")

    add_heading(doc, "阶段 1：关闭 P0 阻断项", 2)
    add_number(doc, "完成 API 级 RBAC、前端权限菜单和最后超级管理员保护。")
    add_number(doc, "接入 Flyway/Liquibase，整理现有 SQL 为正式迁移版本。")
    add_number(doc, "完成网页渠道安全入口，或从首发范围中移除网页渠道。")
    add_number(doc, "接通渠道配置、意图、回复策略；无法接通的功能暂时下线。")
    add_number(doc, "配置 Embedding 模型，修正知识库处理状态并完成 RAG 端到端验收。")
    add_number(doc, "完成模型密钥加密、日志脱敏和现有密钥轮换。")

    add_heading(doc, "阶段 2：建立生产保障", 2)
    add_number(doc, "补齐 Nginx 上传限制、HTTPS、安全头、CORS 和网络隔离。")
    add_number(doc, "为文档处理引入持久任务机制，为外部调用配置真实超时、重试和熔断。")
    add_number(doc, "完成真实转人工流程、客户建档、工单建单和后台操作审计。")
    add_number(doc, "增加 liveness/readiness、指标、结构化日志、告警和费用监控。")

    add_heading(doc, "阶段 3：发布工程化", 2)
    add_number(doc, "Admin 模块补集成测试，前端补 lint、单元测试和关键路径 E2E。")
    add_number(doc, "CI 中强制运行迁移校验、测试、依赖审计、镜像扫描和制品签名。")
    add_number(doc, "冻结发布分支和版本号，整理变更记录、回滚方案与数据备份。")
    add_number(doc, "完成灰度、压测、故障演练和上线评审后再开放真实流量。")

    add_heading(doc, "建议首发范围", 1)
    add_status_table(
        doc,
        ["功能", "建议", "进入首发的前置条件"],
        [
            ["管理员后台", "保留", "完成 RBAC、管理员保护、审计日志"],
            ["AI 试聊", "保留", "禁用成功型 Mock 兜底或明确标识降级"],
            ["FAQ 知识库", "保留", "完成命中率样本验收与变更审计"],
            ["文档 RAG", "条件保留", "Embedding、状态机、审核和检索全部通过"],
            ["钉钉/企微", "条件保留", "真实平台回调、重试、签名和告警通过"],
            ["网页客服", "暂缓", "完成安全入口和真实组件端到端测试后再启用"],
            ["意图/回复策略", "暂缓", "接入消息处理主链路后再展示"],
            ["客户/工单/操作日志", "暂缓", "完成自动建档、建单和审计闭环"],
        ],
        [2200, 1500, 5180],
        status_col=1,
    )

    add_heading(doc, "生产上线验收清单", 1)
    checklist = [
        "所有 P0 问题关闭，并有对应自动化测试或可复现验收记录。",
        "角色权限矩阵经产品、研发和安全共同确认，越权测试全部通过。",
        "数据库从基线版本升级、重复启动和失败回滚均验证通过。",
        "模型密钥、JWT、数据库、Redis、MinIO 和渠道密钥均完成轮换与权限收敛。",
        "知识库完成上传、解析、向量化、审核、检索、引用、删除全流程。",
        "真实钉钉/企微消息完成签名校验、去重、超时、失败重试与告警验证。",
        "HTTPS、安全头、CORS、限流、上传限制和网络隔离均按生产域名生效。",
        "关键接口完成容量与稳定性压测，外部模型异常不会拖垮应用线程。",
        "备份与恢复演练成功，具备明确版本、变更单、回滚步骤和负责人。",
        "灰度环境连续稳定运行并完成业务方验收后，再逐步开放生产流量。",
    ]
    for item in checklist:
        add_bullet(doc, item)

    add_heading(doc, "审计证据索引", 1)
    add_body(doc, "以下路径用于研发定位。行号以 2026-07-13 审计时工作区为准，后续代码修改可能导致偏移。", after=8)
    evidence_rows = [
        ["权限边界", "feisheng-bot-admin/.../config/SecurityConfig.java", "28-40"],
        ["用户删除与角色分配", "feisheng-bot-admin/.../controller/UserController.java", "42, 52-59"],
        ["文档上传与异步处理", "feisheng-bot-admin/.../controller/AdminDocController.java", "84-104, 152-168"],
        ["AI 密钥实体", "feisheng-bot-admin/.../entity/BotAiModelConfig.java", "14"],
        ["未接入的渠道配置", "feisheng-bot-admin/.../controller/ChannelConfigController.java", "18-21"],
        ["未接入的意图管理", "feisheng-bot-admin/.../controller/IntentController.java", "35-40"],
        ["未接入的回复策略", "feisheng-bot-admin/.../controller/ReplyStrategyController.java", "27-33"],
        ["数据库迁移脚本", "feisheng-bot-parent/sql/05_add_minio_storage.sql", "5-12"],
        ["Nginx 代理范围", "docker/nginx/nginx.conf", "9-27"],
        ["生产构建跳过测试", "Dockerfile", "7"],
        ["固定健康检查", "feisheng-bot-admin/.../controller/HealthController.java", "12"],
        ["宽松 CORS", "feisheng-bot-admin/.../config/CorsConfig.java", "12-15"],
    ]
    add_status_table(doc, ["主题", "文件", "行号"], evidence_rows, [2100, 5530, 1250], status_col=None)

    add_callout(
        doc,
        "最终建议",
        "完成阶段 1 后可进入受控灰度；完成阶段 2、测试与故障演练后，才建议按生产标准正式开放。",
        "warn",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
